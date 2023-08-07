import os
import openpyxl
import shutil
from docxtpl import DocxTemplate
from docx.shared import Inches, Cm, Pt
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

script_dir = os.path.dirname(os.path.abspath(__file__))

input_folder = os.path.join(script_dir, 'inputs')
files_in_folder = os.listdir(input_folder)

# Path to the Avery label template in the "template" folder
template_folder = os.path.join(script_dir, 'template_ignore_me')
template_path = os.path.join(template_folder, 'AveryTemplate.docx')  # Replace with your template file name

# Load the template
doc = DocxTemplate(template_path)

excel_file = next(file for file in files_in_folder if file.endswith('.xls') or file.endswith('.xlsx'))
excel_path = os.path.join(input_folder, excel_file)

workbook = openpyxl.load_workbook(excel_path)
sheet = workbook.active

data = []

for row in sheet.iter_rows(min_row=2, values_only=True):
    data.append(row)

workbook.close()

columns = list(zip(*data))

old_ports = [port for port in columns[0] if port is not None]
new_ports = [port for port in columns[1] if port is not None]

if len(old_ports) != len(new_ports):
    print("OLD PORT AND NEW PORTS ARE NOT THE SAME LENGTH. Ensure that the number of new ports lines up with the number of new ports.")
    exit(1)
    
# Stripping Gi and any /0/
for i in range(len(old_ports)):
    if "/0/" in old_ports[i]:
        old_ports[i] = old_ports[i].replace("/0/", "/")
    if old_ports[i].startswith("Gi"):
        old_ports[i] = old_ports[i][2:]
    if old_ports[i].startswith('0'):
        old_ports[i] = old_ports[i][1:]
# I KNOW THIS IS BAD CODING BUT SILENCE
    if "/0/" in new_ports[i]:
        new_ports[i] = new_ports[i].replace("/0/", "/")
    if new_ports[i].startswith("Gi"):
        new_ports[i] = new_ports[i][2:]
    if new_ports[i].startswith('0'):
        old_ports[i] = new_ports[i][1:]
    

print("Old Ports:", old_ports)
print("New Ports:", new_ports)

switch_name = os.path.splitext(excel_file)[0]

# Define the output folder path
output_folder = os.path.join(script_dir, 'output')
os.makedirs(output_folder, exist_ok=True)

# Define the output path for the populated document
output_name = switch_name + '.docx'
output_path = os.path.join(output_folder, output_name)

# Clone the template to the output folder
os.makedirs(os.path.dirname(output_path), exist_ok=True)
shutil.copy2(template_path, output_path)


document = docx.Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.43)
    section.left_margin = Inches(.62)
    section.right_margin = Inches(.31)



# Calculate the maximum number of rows per page
max_rows_per_page = 15

# Calculate the number of pages needed based on the data
num_rows = len(old_ports)
num_pages = (num_rows + max_rows_per_page - 1) // max_rows_per_page

# Add a table with the calculated number of rows and two columns to split the page
num_cols = 2

index = 0  # Track the current index in the old_ports list

for page in range(num_pages):
    table = document.add_table(rows=max_rows_per_page, cols=num_cols)

    # Set cell properties and add content to each cell
    for row in table.rows:
        if num_rows <= 0:
            break  # Break loop if no more rows to add
        
        for cell in row.cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center alignment
            if index < len(old_ports):
                run = paragraph.add_run(old_ports[index]+ "   |    " + new_ports[index])  # Add content from your data list
                run.font.size = Pt(30)  # Adjust font size as needed
                run.font.name = 'Stratum 2 MD'  # Use the desired font
            
            num_rows -= 1
            index += 1

document.save(output_path)

print(f"Populated labels saved at: {output_path}")


# # Add a paragraph with the text 'Hello'
# p1 = document.add_paragraph()

# # Add a run to the first paragraph and set its font size
# run = p1.add_run()
# font = run.font
# font.size = Pt(30)
# font.name = 'Stratum2 MD'

# for i in range(len(old_ports)):
#     # Prints old port plus spaces and new port
#     run = p1.add_run(old_ports[i]+ "         |           " + new_ports[i])
#     font = run.font
#     font.size = Pt(30)
#     font.name = 'Stratum2 MD'
#     if i % 2:
#         print("I is odd")
#         run = p1.add_run("     ")
        
#         # i is odd, add space
#     else:
#         run = p1.add_run("\n")
        

#     print("i: " + str(i))